"""Create a test contract with potentially non-compliant clauses."""

from docx import Document
from pathlib import Path

# Create a new document
doc = Document()

# Add title
doc.add_heading('SERVICE AGREEMENT', 0)

doc.add_paragraph('This Service Agreement ("Agreement") is entered into as of November 10, 2025.')

# Add clauses that might be non-compliant
clauses = [
    {
        'title': '1. LIMITATION OF LIABILITY',
        'text': 'Under no circumstances shall either party be liable for any indirect, incidental, consequential, special or exemplary damages arising out of or related to this Agreement, even if such party has been advised of the possibility of such damages. The total liability of either party shall not exceed $100.'
    },
    {
        'title': '2. INTELLECTUAL PROPERTY',
        'text': 'All intellectual property created during the performance of services under this Agreement shall become the exclusive property of the Client immediately upon creation, including any pre-existing intellectual property of the Service Provider that is incorporated into the deliverables.'
    },
    {
        'title': '3. PAYMENT TERMS',
        'text': 'Payment shall be due within 120 days of invoice date. Late payments will incur a penalty of 5% per month, compounded daily.'
    },
    {
        'title': '4. TERMINATION',
        'text': 'This Agreement may be terminated by the Client at any time without cause and without any payment obligations for work already performed. The Service Provider may not terminate this Agreement under any circumstances.'
    },
    {
        'title': '5. CONFIDENTIALITY',
        'text': 'All information shared by either party shall be considered confidential for a period of 10 years from the date of disclosure. This includes publicly available information.'
    },
    {
        'title': '6. INDEMNIFICATION',
        'text': 'The Service Provider shall indemnify, defend, and hold harmless the Client from any and all claims, damages, losses, and expenses, including but not limited to legal fees, arising from the Service Provider\'s performance under this Agreement, regardless of fault.'
    },
    {
        'title': '7. GOVERNING LAW',
        'text': 'This Agreement shall be governed by the laws of the Cayman Islands, and any disputes shall be resolved exclusively through arbitration in the Cayman Islands, with proceedings conducted in Mandarin Chinese.'
    },
    {
        'title': '8. DATA PROTECTION',
        'text': 'The Service Provider may collect, store, and use any data obtained during the performance of services for its own commercial purposes without restriction or notification to the Client.'
    },
    {
        'title': '9. WARRANTY DISCLAIMER',
        'text': 'THE SERVICE PROVIDER PROVIDES ALL SERVICES "AS IS" WITHOUT ANY WARRANTIES OF ANY KIND, EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE, OR NON-INFRINGEMENT.'
    },
    {
        'title': '10. NON-COMPETE',
        'text': 'The Service Provider agrees not to engage in any business activities that compete with the Client\'s business for a period of 5 years following termination of this Agreement, globally.'
    }
]

# Add each clause
for clause in clauses:
    doc.add_heading(clause['title'], level=1)
    doc.add_paragraph(clause['text'])
    doc.add_paragraph()  # Add spacing

# Save the document
output_path = Path('./sample_test_contract.docx')
doc.save(output_path)
print(f"Test contract created: {output_path.absolute()}")