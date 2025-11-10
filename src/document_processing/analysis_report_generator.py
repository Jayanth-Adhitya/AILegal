"""Generate detailed analysis reports showing all clauses checked and changed."""

import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


class AnalysisReportGenerator:
    """Generate comprehensive analysis reports for transparency."""

    def __init__(self):
        """Initialize report generator."""
        logger.info("Initialized AnalysisReportGenerator")

    def generate_detailed_report(
        self,
        analysis_results: List[Dict[str, Any]],
        summary: Dict[str, Any],
        contract_info: Dict[str, Any],
        output_path: str
    ):
        """
        Generate a detailed analysis report document.

        Args:
            analysis_results: List of clause analysis results
            summary: Contract summary
            contract_info: Contract metadata
            output_path: Path to save the report
        """
        logger.info("Generating detailed analysis report")

        doc = Document()

        # Set up styles
        self._setup_styles(doc)

        # Title page
        self._add_title_page(doc, contract_info)

        # Executive summary
        self._add_executive_summary(doc, summary, analysis_results)

        # Statistics dashboard
        self._add_statistics_dashboard(doc, analysis_results)

        # Detailed clause-by-clause analysis
        self._add_detailed_analysis(doc, analysis_results)

        # Risk assessment matrix
        self._add_risk_matrix(doc, analysis_results)

        # Recommendations
        self._add_recommendations(doc, analysis_results, summary)

        # Save report
        doc.save(output_path)
        logger.info(f"Detailed report saved to: {output_path}")

    def _setup_styles(self, doc: Document):
        """Set up custom styles for the report."""
        try:
            styles = doc.styles

            # Title style
            if 'Report Title' not in styles:
                title_style = styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
                title_style.font.name = 'Arial'
                title_style.font.size = Pt(24)
                title_style.font.bold = True
                title_style.font.color.rgb = RGBColor(0, 51, 102)

            # Section heading style
            if 'Report Section' not in styles:
                section_style = styles.add_style('Report Section', WD_STYLE_TYPE.PARAGRAPH)
                section_style.font.name = 'Arial'
                section_style.font.size = Pt(16)
                section_style.font.bold = True
                section_style.font.color.rgb = RGBColor(0, 102, 204)

            # Clause heading style
            if 'Clause Header' not in styles:
                clause_style = styles.add_style('Clause Header', WD_STYLE_TYPE.PARAGRAPH)
                clause_style.font.name = 'Arial'
                clause_style.font.size = Pt(12)
                clause_style.font.bold = True
                clause_style.font.color.rgb = RGBColor(51, 51, 51)

        except Exception as e:
            logger.warning(f"Could not create custom styles: {e}")

    def _add_title_page(self, doc: Document, contract_info: Dict[str, Any]):
        """Add title page to report."""
        # Title
        title = doc.add_paragraph("AI Legal Assistant")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            title.style = 'Report Title'
        except:
            title.runs[0].font.size = Pt(24)
            title.runs[0].font.bold = True

        # Subtitle
        subtitle = doc.add_paragraph("Detailed Contract Analysis Report")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)

        doc.add_paragraph()  # Spacing

        # Contract info
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"Contract: {contract_info.get('title', 'Untitled')}\n").bold = True
        info_para.add_run(f"File: {contract_info.get('file_name', 'N/A')}\n")
        info_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        info_para.add_run(f"Powered by: Gemini 2.5 Flash (Batch Mode)")

        doc.add_page_break()

    def _add_executive_summary(
        self,
        doc: Document,
        summary: Dict[str, Any],
        analysis_results: List[Dict[str, Any]]
    ):
        """Add executive summary section."""
        heading = doc.add_paragraph("Executive Summary")
        try:
            heading.style = 'Report Section'
        except:
            heading.runs[0].font.size = Pt(16)
            heading.runs[0].font.bold = True

        doc.add_paragraph()

        # Key findings
        findings = doc.add_paragraph()
        findings.add_run("Contract Type: ").bold = True
        findings.add_run(f"{summary.get('contract_type', 'Unknown')}\n")

        findings.add_run("Overall Risk: ").bold = True
        risk = summary.get('overall_risk_assessment', 'unknown').upper()
        risk_run = findings.add_run(f"{risk}\n")
        risk_run.bold = True

        # Color code risk
        risk_colors = {
            "LOW": RGBColor(0, 128, 0),
            "MEDIUM": RGBColor(255, 165, 0),
            "HIGH": RGBColor(255, 69, 0),
            "CRITICAL": RGBColor(178, 34, 34)
        }
        if risk in risk_colors:
            risk_run.font.color.rgb = risk_colors[risk]

        findings.add_run("Total Clauses Reviewed: ").bold = True
        findings.add_run(f"{summary.get('total_clauses_reviewed', len(analysis_results))}\n")

        findings.add_run("Compliant Clauses: ").bold = True
        compliant = summary.get('compliant_clauses', 0)
        findings.add_run(f"{compliant} ").font.color.rgb = RGBColor(0, 128, 0)

        findings.add_run("Non-Compliant: ").bold = True
        non_compliant = summary.get('non_compliant_clauses', 0)
        findings.add_run(f"{non_compliant}\n").font.color.rgb = RGBColor(220, 20, 60)

        # Executive summary text
        if summary.get('executive_summary'):
            doc.add_paragraph()
            doc.add_paragraph(summary['executive_summary'])

        doc.add_paragraph()

    def _add_statistics_dashboard(
        self,
        doc: Document,
        analysis_results: List[Dict[str, Any]]
    ):
        """Add statistics dashboard with visualizations."""
        heading = doc.add_paragraph("Analysis Statistics")
        try:
            heading.style = 'Report Section'
        except:
            heading.runs[0].font.size = Pt(16)
            heading.runs[0].font.bold = True

        # Count by clause type
        clause_types = {}
        risk_levels = {"low": 0, "medium": 0, "high": 0, "critical": 0}
        compliant_count = 0
        non_compliant_count = 0

        for result in analysis_results:
            # Clause types
            clause_type = result.get('clause_type', 'unknown')
            clause_types[clause_type] = clause_types.get(clause_type, 0) + 1

            # Risk levels
            risk = result.get('risk_level', 'unknown').lower()
            if risk in risk_levels:
                risk_levels[risk] += 1

            # Compliance
            if result.get('compliant'):
                compliant_count += 1
            else:
                non_compliant_count += 1

        # Create table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = "Metric"
        header_cells[1].text = "Count"

        # Add rows
        table.add_row().cells[0].text = "Total Clauses"
        table.rows[-1].cells[1].text = str(len(analysis_results))

        table.add_row().cells[0].text = "Compliant"
        table.rows[-1].cells[1].text = f"{compliant_count} ({compliant_count/len(analysis_results)*100:.1f}%)"

        table.add_row().cells[0].text = "Non-Compliant"
        table.rows[-1].cells[1].text = f"{non_compliant_count} ({non_compliant_count/len(analysis_results)*100:.1f}%)"

        doc.add_paragraph()

        # Clause types breakdown
        doc.add_paragraph("Clause Types Distribution:").runs[0].bold = True
        for clause_type, count in sorted(clause_types.items(), key=lambda x: -x[1]):
            doc.add_paragraph(f"  • {clause_type.replace('_', ' ').title()}: {count}", style='List Bullet')

        doc.add_paragraph()

        # Risk levels
        doc.add_paragraph("Risk Levels Distribution:").runs[0].bold = True
        for risk, count in risk_levels.items():
            if count > 0:
                para = doc.add_paragraph(f"  • {risk.upper()}: {count}", style='List Bullet')

        doc.add_page_break()

    def _add_detailed_analysis(
        self,
        doc: Document,
        analysis_results: List[Dict[str, Any]]
    ):
        """Add detailed clause-by-clause analysis."""
        heading = doc.add_paragraph("Detailed Clause Analysis")
        try:
            heading.style = 'Report Section'
        except:
            heading.runs[0].font.size = Pt(16)
            heading.runs[0].font.bold = True

        doc.add_paragraph(
            "This section provides a comprehensive analysis of each clause in the contract, "
            "including compliance status, identified issues, and recommended changes."
        )

        doc.add_paragraph()

        for i, result in enumerate(analysis_results, 1):
            self._add_clause_detail(doc, i, result)

            # Add page break every 5 clauses for readability
            if i % 5 == 0 and i < len(analysis_results):
                doc.add_page_break()

    def _add_clause_detail(
        self,
        doc: Document,
        clause_num: int,
        result: Dict[str, Any]
    ):
        """Add detailed information for a single clause."""
        # Clause header with status indicator
        compliant = result.get('compliant', None)
        status_emoji = "✅" if compliant else "❌" if compliant is False else "⚠️"

        header = doc.add_paragraph(f"\n{status_emoji} Clause {clause_num}: {result.get('clause_type', 'Unknown').replace('_', ' ').title()}")
        try:
            header.style = 'Clause Header'
        except:
            header.runs[0].font.size = Pt(12)
            header.runs[0].font.bold = True

        # Original text
        doc.add_paragraph("Original Text:").runs[0].bold = True
        original = doc.add_paragraph(result.get('text', 'N/A'))
        original.paragraph_format.left_indent = Inches(0.25)
        original.runs[0].font.italic = True
        original.runs[0].font.size = Pt(10)

        # Analysis details
        details = doc.add_paragraph()

        details.add_run("Clause Type: ").bold = True
        details.add_run(f"{result.get('clause_type', 'Unknown').replace('_', ' ').title()}\n")

        details.add_run("Compliance Status: ").bold = True
        if compliant:
            status_run = details.add_run("COMPLIANT ✓\n")
            status_run.font.color.rgb = RGBColor(0, 128, 0)
            status_run.bold = True
        elif compliant is False:
            status_run = details.add_run("NON-COMPLIANT ✗\n")
            status_run.font.color.rgb = RGBColor(220, 20, 60)
            status_run.bold = True
        else:
            details.add_run("REQUIRES REVIEW ⚠\n")

        details.add_run("Risk Level: ").bold = True
        risk = result.get('risk_level', 'Unknown').upper()
        risk_run = details.add_run(f"{risk}\n")
        risk_colors = {
            "LOW": RGBColor(0, 128, 0),
            "MEDIUM": RGBColor(255, 165, 0),
            "HIGH": RGBColor(255, 69, 0),
            "CRITICAL": RGBColor(178, 34, 34)
        }
        if risk in risk_colors:
            risk_run.font.color.rgb = risk_colors[risk]

        # Issues found (if non-compliant)
        if not compliant and result.get('issues'):
            doc.add_paragraph("Issues Identified:").runs[0].bold = True

            for issue in result['issues']:
                issue_para = doc.add_paragraph(style='List Bullet')
                # Issues are now simple strings, not nested objects
                if isinstance(issue, str):
                    issue_para.add_run(issue)
                else:
                    # Fallback for old format
                    issue_para.add_run(f"{issue.get('issue_description', str(issue))}\n")
                    if issue.get('policy_reference'):
                        issue_para.add_run(f"  Policy: {issue['policy_reference']}\n").font.italic = True

        # Recommendations (new field)
        if result.get('recommendations'):
            doc.add_paragraph("Recommendations:").runs[0].bold = True
            for rec in result['recommendations']:
                rec_para = doc.add_paragraph(style='List Bullet')
                # Recommendations are now simple strings
                if isinstance(rec, str):
                    rec_para.add_run(rec)
                else:
                    rec_para.add_run(str(rec))

        # Suggested changes (if non-compliant)
        if not compliant and result.get('suggested_alternative'):
            doc.add_paragraph("Recommended Alternative:").runs[0].bold = True
            suggestion_para = doc.add_paragraph(result['suggested_alternative'])
            suggestion_para.paragraph_format.left_indent = Inches(0.25)
            suggestion_para.runs[0].font.color.rgb = RGBColor(0, 100, 0)
            suggestion_para.runs[0].font.bold = True

        # Policy citations
        if result.get('policy_references'):
            doc.add_paragraph("Policy References:").runs[0].bold = True
            for citation in result['policy_references']:
                # Policy references are now simple strings
                if isinstance(citation, str):
                    doc.add_paragraph(f"  • {citation}", style='List Bullet')
                else:
                    doc.add_paragraph(f"  • {str(citation)}", style='List Bullet')

        # Separator
        doc.add_paragraph("─" * 80)

    def _add_risk_matrix(
        self,
        doc: Document,
        analysis_results: List[Dict[str, Any]]
    ):
        """Add risk assessment matrix."""
        doc.add_page_break()

        heading = doc.add_paragraph("Risk Assessment Matrix")
        try:
            heading.style = 'Report Section'
        except:
            heading.runs[0].font.size = Pt(16)
            heading.runs[0].font.bold = True

        # Group by risk and clause type
        risk_matrix = {}
        for result in analysis_results:
            clause_type = result.get('clause_type', 'unknown')
            risk = result.get('risk_level', 'unknown')

            if risk not in risk_matrix:
                risk_matrix[risk] = []

            if not result.get('compliant'):
                risk_matrix[risk].append({
                    'type': clause_type,
                    'text': result.get('text', '')[:100] + '...'
                })

        # Display by risk level
        for risk_level in ['critical', 'high', 'medium', 'low']:
            if risk_level in risk_matrix and risk_matrix[risk_level]:
                risk_heading = doc.add_paragraph(f"\n{risk_level.upper()} RISK CLAUSES:")
                risk_heading.runs[0].bold = True

                risk_colors = {
                    "critical": RGBColor(178, 34, 34),
                    "high": RGBColor(255, 69, 0),
                    "medium": RGBColor(255, 165, 0),
                    "low": RGBColor(255, 215, 0)
                }
                if risk_level in risk_colors:
                    risk_heading.runs[0].font.color.rgb = risk_colors[risk_level]

                for item in risk_matrix[risk_level]:
                    clause_para = doc.add_paragraph(style='List Bullet')
                    clause_para.add_run(f"{item['type'].replace('_', ' ').title()}: ").bold = True
                    clause_para.add_run(f"{item['text']}")

    def _add_recommendations(
        self,
        doc: Document,
        analysis_results: List[Dict[str, Any]],
        summary: Dict[str, Any]
    ):
        """Add recommendations section."""
        doc.add_page_break()

        heading = doc.add_paragraph("Recommendations & Next Steps")
        try:
            heading.style = 'Report Section'
        except:
            heading.runs[0].font.size = Pt(16)
            heading.runs[0].font.bold = True

        # Overall recommendation
        if summary.get('recommendation'):
            doc.add_paragraph("Overall Assessment:").runs[0].bold = True
            doc.add_paragraph(summary['recommendation'])

        doc.add_paragraph()

        # Critical actions
        doc.add_paragraph("Immediate Actions Required:").runs[0].bold = True

        critical_issues = [r for r in analysis_results if r.get('risk_level') == 'critical']
        if critical_issues:
            for issue in critical_issues:
                action = doc.add_paragraph(style='List Number')
                action.add_run(f"Address {issue.get('clause_type', 'unknown').replace('_', ' ')} clause: ").bold = True
                action.add_run(issue.get('rejection_reason', 'High risk identified'))
        else:
            doc.add_paragraph("  • No critical issues identified", style='List Bullet')

        doc.add_paragraph()

        # SME review needed
        if summary.get('requires_sme_review'):
            doc.add_paragraph("Subject Matter Expert Review Needed For:").runs[0].bold = True
            for area in summary['requires_sme_review']:
                doc.add_paragraph(f"  • {area}", style='List Bullet')

        doc.add_paragraph()

        # Footer
        doc.add_paragraph("─" * 80)
        footer = doc.add_paragraph(
            "\nThis report was generated by AI Legal Assistant using Gemini 2.5 Flash. "
            "All recommendations should be reviewed by qualified legal counsel before implementation."
        )
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.italic = True

    def generate_changes_summary_html(
        self,
        analysis_results: List[Dict[str, Any]],
        output_path: str
    ):
        """
        Generate an HTML summary of all changes for web viewing.

        Args:
            analysis_results: List of clause analysis results
            output_path: Path to save HTML file
        """
        logger.info("Generating HTML changes summary")

        html_content = self._build_html_summary(analysis_results)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        logger.info(f"HTML summary saved to: {output_path}")

    def _build_html_summary(self, analysis_results: List[Dict[str, Any]]) -> str:
        """Build HTML content for changes summary."""
        html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Contract Analysis Summary</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        .header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 30px;
            border-radius: 10px;
            margin-bottom: 30px;
        }}
        .stats {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin-bottom: 30px;
        }}
        .stat-card {{
            background: white;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            text-align: center;
        }}
        .stat-value {{
            font-size: 2em;
            font-weight: bold;
            margin: 10px 0;
        }}
        .compliant {{ color: #10b981; }}
        .non-compliant {{ color: #ef4444; }}
        .clause {{
            background: white;
            padding: 20px;
            margin-bottom: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        .clause-header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 15px;
            padding-bottom: 10px;
            border-bottom: 2px solid #e5e7eb;
        }}
        .clause-status {{
            padding: 5px 15px;
            border-radius: 20px;
            font-weight: bold;
            font-size: 0.9em;
        }}
        .status-compliant {{
            background: #d1fae5;
            color: #065f46;
        }}
        .status-non-compliant {{
            background: #fee2e2;
            color: #991b1b;
        }}
        .risk-badge {{
            display: inline-block;
            padding: 3px 10px;
            border-radius: 12px;
            font-size: 0.85em;
            font-weight: bold;
            margin-left: 10px;
        }}
        .risk-critical {{ background: #fecaca; color: #991b1b; }}
        .risk-high {{ background: #fed7aa; color: #9a3412; }}
        .risk-medium {{ background: #fde68a; color: #78350f; }}
        .risk-low {{ background: #d9f99d; color: #365314; }}
        .original-text {{
            background: #f9fafb;
            padding: 15px;
            border-left: 4px solid #9ca3af;
            margin: 15px 0;
            font-style: italic;
        }}
        .suggestion {{
            background: #d1fae5;
            padding: 15px;
            border-left: 4px solid #10b981;
            margin: 15px 0;
        }}
        .issues {{
            background: #fee2e2;
            padding: 15px;
            border-left: 4px solid #ef4444;
            margin: 15px 0;
        }}
        .policy-ref {{
            display: inline-block;
            background: #e0e7ff;
            color: #3730a3;
            padding: 3px 8px;
            border-radius: 4px;
            font-size: 0.85em;
            margin: 2px;
        }}
        .filter-buttons {{
            margin: 20px 0;
            display: flex;
            gap: 10px;
            flex-wrap: wrap;
        }}
        .filter-btn {{
            padding: 8px 16px;
            border: none;
            border-radius: 6px;
            cursor: pointer;
            font-weight: 500;
            transition: all 0.3s;
        }}
        .filter-btn.active {{
            transform: scale(1.05);
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>🤖 AI Legal Assistant - Analysis Summary</h1>
        <p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>

    <div class="stats">
"""

        # Calculate statistics
        total = len(analysis_results)
        compliant = sum(1 for r in analysis_results if r.get('compliant'))
        non_compliant = total - compliant

        html += f"""
        <div class="stat-card">
            <div class="stat-label">Total Clauses</div>
            <div class="stat-value">{total}</div>
        </div>
        <div class="stat-card">
            <div class="stat-label">Compliant</div>
            <div class="stat-value compliant">{compliant}</div>
        </div>
        <div class="stat-card">
            <div class="stat-label">Non-Compliant</div>
            <div class="stat-value non-compliant">{non_compliant}</div>
        </div>
        <div class="stat-card">
            <div class="stat-label">Compliance Rate</div>
            <div class="stat-value">{(compliant/total*100):.1f}%</div>
        </div>
    </div>

    <div class="filter-buttons">
        <button class="filter-btn active" onclick="filterClauses('all')" style="background: #667eea; color: white;">All Clauses</button>
        <button class="filter-btn" onclick="filterClauses('non-compliant')" style="background: #ef4444; color: white;">Non-Compliant Only</button>
        <button class="filter-btn" onclick="filterClauses('critical')" style="background: #991b1b; color: white;">Critical Risk</button>
    </div>

    <div id="clauses-container">
"""

        # Add each clause
        for i, result in enumerate(analysis_results, 1):
            compliant = result.get('compliant')
            risk = result.get('risk_level', 'unknown').lower()

            status_class = 'status-compliant' if compliant else 'status-non-compliant'
            status_text = '✅ COMPLIANT' if compliant else '❌ NON-COMPLIANT'
            data_attrs = f'data-compliant="{str(compliant).lower()}" data-risk="{risk}"'

            html += f"""
        <div class="clause" {data_attrs}>
            <div class="clause-header">
                <h3>Clause {i}: {result.get('clause_type', 'Unknown').replace('_', ' ').title()}</h3>
                <div>
                    <span class="clause-status {status_class}">{status_text}</span>
                    <span class="risk-badge risk-{risk}">{risk.upper()} RISK</span>
                </div>
            </div>

            <div class="original-text">
                <strong>Original Text:</strong><br>
                {result.get('text', 'N/A')[:500]}{'...' if len(result.get('text', '')) > 500 else ''}
            </div>
"""

            if not compliant and result.get('issues'):
                html += '<div class="issues"><strong>Issues Identified:</strong><ul>'
                for issue in result['issues']:
                    html += f"<li>{issue.get('issue_description', 'Unknown')}<br>"
                    html += f"<small>Policy: {issue.get('policy_reference', 'N/A')}</small></li>"
                html += '</ul></div>'

            if not compliant and result.get('redline_suggestion'):
                html += f"""
            <div class="suggestion">
                <strong>💡 Recommended Alternative:</strong><br>
                {result['redline_suggestion']}
            </div>
"""

            if result.get('policy_citations'):
                html += '<div><strong>Policy References:</strong><br>'
                for citation in result['policy_citations']:
                    html += f'<span class="policy-ref">{citation}</span>'
                html += '</div>'

            html += '</div>'

        html += """
    </div>

    <script>
        function filterClauses(filter) {
            const clauses = document.querySelectorAll('.clause');
            const buttons = document.querySelectorAll('.filter-btn');

            buttons.forEach(btn => btn.classList.remove('active'));
            event.target.classList.add('active');

            clauses.forEach(clause => {
                if (filter === 'all') {
                    clause.style.display = 'block';
                } else if (filter === 'non-compliant') {
                    clause.style.display = clause.dataset.compliant === 'false' ? 'block' : 'none';
                } else if (filter === 'critical') {
                    clause.style.display = clause.dataset.risk === 'critical' ? 'block' : 'none';
                }
            });
        }
    </script>
</body>
</html>
"""

        return html
