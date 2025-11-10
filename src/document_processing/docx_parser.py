"""Parse Word documents and extract contract content."""

import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


class DocxParser:
    """Parse Word documents and extract structured content."""

    def __init__(self, file_path: str):
        """
        Initialize parser with a Word document.

        Args:
            file_path: Path to the .docx file
        """
        self.file_path = Path(file_path)
        self.document = Document(str(self.file_path))
        logger.info(f"Loaded document: {self.file_path.name}")

    def extract_paragraphs(self) -> List[Dict[str, Any]]:
        """
        Extract all paragraphs with their properties.

        Returns:
            List of paragraph dictionaries with text and metadata
        """
        paragraphs = []

        for i, para in enumerate(self.document.paragraphs):
            if para.text.strip():  # Skip empty paragraphs
                paragraphs.append({
                    "index": i,
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal",
                    "runs": len(para.runs),
                    "is_heading": self._is_heading(para)
                })

        logger.info(f"Extracted {len(paragraphs)} paragraphs")
        return paragraphs

    def _is_heading(self, paragraph: Paragraph) -> bool:
        """Check if paragraph is a heading."""
        if paragraph.style and paragraph.style.name:
            return "Heading" in paragraph.style.name or "Title" in paragraph.style.name
        return False

    def extract_sections(self) -> List[Dict[str, Any]]:
        """
        Extract document sections based on headings.

        Returns:
            List of sections with headings and content
        """
        sections = []
        current_section = None

        for i, para in enumerate(self.document.paragraphs):
            if self._is_heading(para):
                # Save previous section
                if current_section:
                    sections.append(current_section)

                # Start new section
                current_section = {
                    "heading": para.text,
                    "heading_style": para.style.name,
                    "heading_index": i,
                    "content": []
                }
            elif current_section and para.text.strip():
                # Add content to current section
                current_section["content"].append({
                    "index": i,
                    "text": para.text
                })

        # Add final section
        if current_section:
            sections.append(current_section)

        logger.info(f"Extracted {len(sections)} sections")
        return sections

    def extract_full_text(self) -> str:
        """
        Extract full document text.

        Returns:
            Complete document text
        """
        full_text = "\n\n".join([para.text for para in self.document.paragraphs if para.text.strip()])
        logger.info(f"Extracted full text: {len(full_text)} characters")
        return full_text

    def extract_tables(self) -> List[Dict[str, Any]]:
        """
        Extract all tables from the document.

        Returns:
            List of table data
        """
        tables_data = []

        for i, table in enumerate(self.document.tables):
            table_data = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "data": []
            }

            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data["data"].append(row_data)

            tables_data.append(table_data)

        logger.info(f"Extracted {len(tables_data)} tables")
        return tables_data

    def extract_comments(self) -> List[Dict[str, Any]]:
        """
        Extract existing comments from the document.

        Returns:
            List of comments with metadata
        """
        comments = []

        # Check if document has comments part
        if hasattr(self.document, 'part') and hasattr(self.document.part, 'comments'):
            try:
                comments_part = self.document.part.comments
                if comments_part and hasattr(comments_part, 'comments'):
                    for comment in comments_part.comments:
                        comments.append({
                            "id": comment.id if hasattr(comment, 'id') else None,
                            "author": comment.author if hasattr(comment, 'author') else "Unknown",
                            "date": str(comment.date) if hasattr(comment, 'date') else None,
                            "text": comment.text if hasattr(comment, 'text') else ""
                        })
            except Exception as e:
                logger.warning(f"Could not extract comments: {e}")

        logger.info(f"Extracted {len(comments)} existing comments")
        return comments

    def get_document_properties(self) -> Dict[str, Any]:
        """
        Get document metadata and properties.

        Returns:
            Document properties dictionary
        """
        core_properties = self.document.core_properties

        properties = {
            "title": core_properties.title or "Untitled",
            "author": core_properties.author or "Unknown",
            "created": str(core_properties.created) if core_properties.created else None,
            "modified": str(core_properties.modified) if core_properties.modified else None,
            "paragraphs_count": len(self.document.paragraphs),
            "tables_count": len(self.document.tables),
            "file_name": self.file_path.name
        }

        logger.info(f"Extracted document properties: {properties['title']}")
        return properties

    def extract_structured_content(self) -> Dict[str, Any]:
        """
        Extract all content in a structured format.

        Returns:
            Complete structured document data
        """
        return {
            "properties": self.get_document_properties(),
            "full_text": self.extract_full_text(),
            "paragraphs": self.extract_paragraphs(),
            "sections": self.extract_sections(),
            "tables": self.extract_tables(),
            "comments": self.extract_comments()
        }

    @staticmethod
    def is_valid_docx(file_path: str) -> bool:
        """
        Check if file is a valid Word document.

        Args:
            file_path: Path to check

        Returns:
            True if valid .docx file
        """
        try:
            path = Path(file_path)
            if not path.exists():
                return False

            if path.suffix.lower() != '.docx':
                return False

            # Try to open it
            Document(str(path))
            return True

        except Exception as e:
            logger.error(f"Invalid DOCX file: {e}")
            return False
