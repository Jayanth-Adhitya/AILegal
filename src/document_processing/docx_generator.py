"""Generate Word documents with track changes and comments."""

import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
from lxml import etree

logger = logging.getLogger(__name__)

# Word Processing XML namespace for OOXML manipulation
WORD_NAMESPACE = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W = '{' + WORD_NAMESPACE + '}'


class DocxGenerator:
    """Generate Word documents with track changes and AI comments."""

    def __init__(self, original_doc_path: Optional[str] = None):
        """
        Initialize generator.

        Args:
            original_doc_path: Path to original document (to preserve formatting)
        """
        if original_doc_path:
            self.document = Document(original_doc_path)
            logger.info(f"Loaded original document: {Path(original_doc_path).name}")
        else:
            self.document = Document()
            logger.info("Created new blank document")

        # Initialize revision counter for track changes
        self._revision_id = 1

    def _generate_revision_id(self) -> str:
        """
        Generate a unique sequential revision ID for track changes.

        Returns:
            String representation of the revision ID
        """
        rev_id = str(self._revision_id)
        self._revision_id += 1
        return rev_id

    def _generate_timestamp(self) -> str:
        """
        Generate an ISO 8601 formatted timestamp for track changes.

        Returns:
            ISO 8601 formatted timestamp string
        """
        return datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ')

    def _add_track_change_deletion(self, paragraph, text: str) -> bool:
        """
        Add a tracked deletion to a paragraph using OOXML markup.

        Creates a <w:del> element containing the text to be deleted,
        which appears as strikethrough in Word's Review pane.

        Args:
            paragraph: python-docx Paragraph object
            text: Text to mark as deleted

        Returns:
            True if successful
        """
        try:
            # Access paragraph XML element
            p_element = paragraph._element

            # Create deletion element with revision metadata
            del_elem = etree.Element(
                f'{W}del',
                attrib={
                    f'{W}id': self._generate_revision_id(),
                    f'{W}author': 'AI Legal Assistant',
                    f'{W}date': self._generate_timestamp()
                }
            )

            # Create run element for deleted text
            run_elem = etree.SubElement(del_elem, f'{W}r')

            # Add run properties to preserve formatting
            rPr = etree.SubElement(run_elem, f'{W}rPr')

            # Add deleted text element
            del_text_elem = etree.SubElement(run_elem, f'{W}delText')
            del_text_elem.text = text

            # Append deletion to paragraph
            p_element.append(del_elem)

            logger.debug(f"Added track change deletion: {text[:50]}...")
            return True

        except Exception as e:
            logger.error(f"Error adding track change deletion: {e}")
            return False

    def _add_track_change_insertion(self, paragraph, text: str) -> bool:
        """
        Add a tracked insertion to a paragraph using OOXML markup.

        Creates a <w:ins> element containing the new text,
        which appears as underlined in Word's Review pane.

        Args:
            paragraph: python-docx Paragraph object
            text: Text to mark as inserted

        Returns:
            True if successful
        """
        try:
            # Access paragraph XML element
            p_element = paragraph._element

            # Create insertion element with revision metadata
            ins_elem = etree.Element(
                f'{W}ins',
                attrib={
                    f'{W}id': self._generate_revision_id(),
                    f'{W}author': 'AI Legal Assistant',
                    f'{W}date': self._generate_timestamp()
                }
            )

            # Create run element for inserted text
            run_elem = etree.SubElement(ins_elem, f'{W}r')

            # Add run properties to preserve formatting
            rPr = etree.SubElement(run_elem, f'{W}rPr')

            # Add text element
            text_elem = etree.SubElement(run_elem, f'{W}t')
            text_elem.text = text

            # Preserve whitespace
            text_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

            # Append insertion to paragraph
            p_element.append(ins_elem)

            logger.debug(f"Added track change insertion: {text[:50]}...")
            return True

        except Exception as e:
            logger.error(f"Error adding track change insertion: {e}")
            return False

    def _enable_track_changes(self) -> bool:
        """
        Enable track changes mode in the document settings.

        This ensures the document opens in Word with the Review pane active
        and all changes visible.

        Returns:
            True if successful
        """
        try:
            # Access document settings
            settings = self.document.settings.element

            # Check if trackRevisions setting already exists
            track_rev = settings.find(f'{W}trackRevisions')

            if track_rev is None:
                # Create trackRevisions element
                track_rev = etree.SubElement(settings, f'{W}trackRevisions')

            # Enable track changes
            track_rev.set(f'{W}val', 'true')

            logger.info("Enabled track changes mode in document settings")
            return True

        except Exception as e:
            logger.error(f"Error enabling track changes: {e}")
            return False

    def add_comment_to_paragraph(
        self,
        paragraph_index: int,
        comment_text: str,
        author: str = "AI Legal Assistant",
        initials: str = "AI"
    ) -> bool:
        """
        Add a comment to a specific paragraph.

        Args:
            paragraph_index: Index of the paragraph to comment on
            comment_text: The comment content
            author: Comment author name
            initials: Author initials

        Returns:
            True if comment added successfully
        """
        try:
            if paragraph_index >= len(self.document.paragraphs):
                logger.error(f"Paragraph index {paragraph_index} out of range")
                return False

            paragraph = self.document.paragraphs[paragraph_index]

            # Add comment to all runs in the paragraph
            if paragraph.runs:
                self.document.add_comment(
                    runs=paragraph.runs,
                    text=comment_text,
                    author=author,
                    initials=initials
                )
                logger.debug(f"Added comment to paragraph {paragraph_index}")
                return True
            else:
                logger.warning(f"Paragraph {paragraph_index} has no runs")
                return False

        except Exception as e:
            logger.error(f"Error adding comment to paragraph {paragraph_index}: {e}")
            return False

    def add_inline_comment(
        self,
        paragraph_index: int,
        comment_text: str,
        highlight_text: bool = True
    ) -> bool:
        """
        Add an inline visual comment (as highlighted text with note).

        Note: python-docx has limitations with Word's native commenting system.
        This method adds a visible indicator in the document.

        Args:
            paragraph_index: Index of the paragraph
            comment_text: The comment content
            highlight_text: Whether to highlight the commented text

        Returns:
            True if successful
        """
        try:
            if paragraph_index >= len(self.document.paragraphs):
                return False

            paragraph = self.document.paragraphs[paragraph_index]

            # Add a visual comment marker
            comment_run = paragraph.add_run(f" [AI COMMENT: {comment_text}]")
            comment_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
            comment_run.font.size = Pt(9)
            comment_run.font.italic = True

            # Highlight original text if requested
            if highlight_text and paragraph.runs:
                for run in paragraph.runs[:-1]:  # Exclude the comment run we just added
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

            logger.debug(f"Added inline comment to paragraph {paragraph_index}")
            return True

        except Exception as e:
            logger.error(f"Error adding inline comment: {e}")
            return False

    def add_redline_suggestion(
        self,
        paragraph_index: int,
        original_text: str,
        suggested_text: str,
        explanation: str
    ) -> bool:
        """
        Add a redline suggestion using native Word track changes.

        Creates tracked deletion for original text and tracked insertion
        for suggested text, which appear in Word's Review pane and can
        be accepted or rejected.

        Args:
            paragraph_index: Index of the paragraph
            original_text: Original clause text to mark as deleted
            suggested_text: Suggested replacement text to mark as inserted
            explanation: Explanation for the change (added as comment)

        Returns:
            True if successful
        """
        try:
            if paragraph_index >= len(self.document.paragraphs):
                logger.error(f"Paragraph index {paragraph_index} out of range")
                return False

            paragraph = self.document.paragraphs[paragraph_index]

            # Remove all existing runs to avoid duplication
            # We must remove the run elements from the paragraph XML, not just clear text
            p_element = paragraph._element
            for run in list(paragraph.runs):  # Create a copy of the list to iterate safely
                p_element.remove(run._element)

            # Add tracked deletion for original text
            self._add_track_change_deletion(paragraph, original_text)

            # Add tracked insertion for suggested text
            self._add_track_change_insertion(paragraph, suggested_text)

            # Add explanation as inline comment (for now)
            # TODO: Future enhancement - use native Word comments
            if explanation:
                self.add_inline_comment(paragraph_index, explanation, highlight_text=False)

            logger.debug(f"Added track changes to paragraph {paragraph_index}")
            return True

        except Exception as e:
            logger.error(f"Error adding track change redline: {e}")
            return False

    def create_review_document(
        self,
        analysis_results: List[Dict[str, Any]],
        output_path: str
    ) -> bool:
        """
        Create a complete review document with native track changes.

        Generates a Word document with tracked deletions and insertions
        that can be reviewed and accepted/rejected in Word's Review pane.

        Args:
            analysis_results: List of clause analysis results
            output_path: Path to save the output document

        Returns:
            True if document created successfully
        """
        try:
            # Enable track changes mode in document settings
            self._enable_track_changes()

            for result in analysis_results:
                paragraph_index = result.get("paragraph_index")
                compliant = result.get("compliant", True)

                if not compliant:
                    # Add rejection comment
                    comment_text = self._format_rejection_comment(result)
                    self.add_inline_comment(paragraph_index, comment_text)

                    # If there's a redline suggestion, add it (support both field names)
                    suggested_text = result.get("suggested_alternative") or result.get("redline_suggestion")
                    if suggested_text:
                        self.add_redline_suggestion(
                            paragraph_index=paragraph_index,
                            original_text=result.get("text", ""),
                            suggested_text=suggested_text,
                            explanation=result.get("rejection_reason", "")
                        )

            # Save the document
            self.document.save(output_path)
            logger.info(f"Saved review document with track changes to: {output_path}")
            return True

        except Exception as e:
            logger.error(f"Error creating review document: {e}")
            return False

    def _format_rejection_comment(self, result: Dict[str, Any]) -> str:
        """Format a rejection comment with policy citations."""
        comment_parts = [
            "REJECTED - Non-compliant clause",
            "",
            f"Reason: {result.get('rejection_reason', 'Policy violation')}",
        ]

        # Add policy citations (support both field names)
        policy_refs = result.get("policy_references") or result.get("policy_citations") or []
        if policy_refs:
            comment_parts.append("")
            comment_parts.append("Policy References:")
            for citation in policy_refs:
                comment_parts.append(f"  - {citation}")

        # Add risk level
        if result.get("risk_level"):
            comment_parts.append("")
            comment_parts.append(f"Risk Level: {result['risk_level'].upper()}")

        # Add suggested action (support both field names)
        suggested_text = result.get("suggested_alternative") or result.get("redline_suggestion")
        if suggested_text:
            comment_parts.append("")
            comment_parts.append("Suggested Alternative:")
            comment_parts.append(suggested_text)

        return "\n".join(comment_parts)

    def add_summary_page(
        self,
        summary: Dict[str, Any]
    ):
        """
        Add a summary page at the beginning of the document.

        Args:
            summary: Contract analysis summary
        """
        # Insert a new paragraph at the beginning
        intro_para = self.document.add_paragraph()
        self.document.paragraphs.insert(0, intro_para)

        # Add title
        title = self.document.add_paragraph("AI Legal Assistant - Contract Review Summary")
        try:
            title.style = "Heading 1"
        except KeyError:
            # If Heading 1 style doesn't exist, format manually
            title_run = title.runs[0] if title.runs else title.add_run()
            title_run.font.size = Pt(16)
            title_run.font.bold = True
        self.document.paragraphs.insert(0, title)

        # Add summary content
        summary_para = self.document.add_paragraph()
        summary_para.add_run("Contract Type: ").bold = True
        summary_para.add_run(f"{summary.get('contract_type', 'Unknown')}\n")

        summary_para.add_run("Total Clauses Reviewed: ").bold = True
        summary_para.add_run(f"{summary.get('total_clauses_reviewed', 0)}\n")

        summary_para.add_run("Compliant Clauses: ").bold = True
        summary_para.add_run(f"{summary.get('compliant_clauses', 0)}\n")

        summary_para.add_run("Non-Compliant Clauses: ").bold = True
        summary_para.add_run(f"{summary.get('non_compliant_clauses', 0)}\n")

        summary_para.add_run("Overall Risk: ").bold = True
        risk_run = summary_para.add_run(f"{summary.get('overall_risk_assessment', 'Unknown').upper()}\n")

        # Color code risk
        risk_colors = {
            "low": RGBColor(0, 128, 0),
            "medium": RGBColor(255, 165, 0),
            "high": RGBColor(255, 0, 0),
            "critical": RGBColor(139, 0, 0)
        }
        risk_level = summary.get('overall_risk_assessment', '').lower()
        if risk_level in risk_colors:
            risk_run.font.color.rgb = risk_colors[risk_level]

        # Add page break
        self.document.add_page_break()

        logger.info("Added summary page to document")

    def save(self, output_path: str):
        """Save the document to a file."""
        self.document.save(output_path)
        logger.info(f"Document saved to: {output_path}")
