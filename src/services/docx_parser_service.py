"""Service for parsing DOCX files and extracting content, formatting, and track changes."""

import os
import uuid
import json
from typing import Dict, List, Any, Optional
from pathlib import Path
from zipfile import ZipFile
from datetime import datetime

from docx import Document as DocxDocument
from lxml import etree


class DocxParserService:
    """Service for parsing Microsoft Word DOCX files."""

    # OOXML namespaces
    NAMESPACES = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    }

    def __init__(self, upload_dir: str = "uploads/documents"):
        """
        Initialize the DOCX parser service.

        Args:
            upload_dir: Base directory for storing uploaded files
        """
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    def save_uploaded_file(self, file_content: bytes, document_id: str, original_filename: str) -> Dict[str, Any]:
        """
        Save uploaded DOCX file to disk.

        Args:
            file_content: Binary content of the DOCX file
            document_id: UUID of the document
            original_filename: Original filename from upload

        Returns:
            Dictionary with file_path, file_name, file_size
        """
        # Create document-specific directory
        doc_dir = self.upload_dir / document_id
        doc_dir.mkdir(parents=True, exist_ok=True)

        # Save file
        file_path = doc_dir / "original.docx"
        with open(file_path, "wb") as f:
            f.write(file_content)

        return {
            "file_path": str(file_path),
            "file_name": original_filename,
            "file_size": len(file_content)
        }

    def parse_docx_structure(self, docx_path: str) -> Dict[str, Any]:
        """
        Parse DOCX file and extract basic structure (paragraphs, tables, formatting).

        Args:
            docx_path: Path to the DOCX file

        Returns:
            Dictionary with parsed content
        """
        try:
            doc = DocxDocument(docx_path)

            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal",
                        "runs": self._extract_runs(para)
                    })

            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            return {
                "success": True,
                "paragraphs": paragraphs,
                "tables": tables,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables)
            }

        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }

    def _extract_runs(self, paragraph) -> List[Dict[str, Any]]:
        """
        Extract formatting runs from a paragraph.

        Args:
            paragraph: python-docx Paragraph object

        Returns:
            List of run dictionaries with text and formatting
        """
        runs = []
        for run in paragraph.runs:
            if run.text:
                runs.append({
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_size": run.font.size.pt if run.font.size else None,
                    "font_name": run.font.name
                })
        return runs

    def extract_track_changes(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        Extract track changes (insertions and deletions) from DOCX file.

        This parses the OOXML directly using lxml to extract <w:ins> and <w:del> elements.

        Args:
            docx_path: Path to the DOCX file

        Returns:
            List of track changes with type, content, author, date, position
        """
        changes = []

        try:
            with ZipFile(docx_path) as docx_zip:
                # Read the main document XML
                xml_content = docx_zip.read('word/document.xml')
                tree = etree.fromstring(xml_content)

                # Extract insertions
                for ins in tree.xpath('.//w:ins', namespaces=self.NAMESPACES):
                    change = self._parse_tracked_change(ins, 'insert')
                    if change:
                        changes.append(change)

                # Extract deletions
                for dels in tree.xpath('.//w:del', namespaces=self.NAMESPACES):
                    change = self._parse_tracked_change(dels, 'delete')
                    if change:
                        changes.append(change)

        except Exception as e:
            print(f"Error extracting track changes: {str(e)}")

        return changes

    def _parse_tracked_change(self, element, change_type: str) -> Optional[Dict[str, Any]]:
        """
        Parse a single tracked change element (<w:ins> or <w:del>).

        Args:
            element: lxml Element representing the change
            change_type: 'insert' or 'delete'

        Returns:
            Dictionary with change details or None if parsing fails
        """
        try:
            # Extract attributes
            author = element.get(f'{{{self.NAMESPACES["w"]}}}author', 'Unknown')
            date_str = element.get(f'{{{self.NAMESPACES["w"]}}}date', '')
            change_id = element.get(f'{{{self.NAMESPACES["w"]}}}id', '')

            # Extract text content
            text_elements = element.xpath('.//w:t/text()', namespaces=self.NAMESPACES)
            content = ''.join(text_elements)

            # Parse date
            try:
                date = datetime.fromisoformat(date_str.replace('Z', '+00:00')) if date_str else datetime.now()
            except:
                date = datetime.now()

            return {
                "type": change_type,
                "content": content,
                "author": author,
                "date": date.isoformat(),
                "change_id": change_id,
                "position": 0  # Will be calculated later based on document structure
            }

        except Exception as e:
            print(f"Error parsing tracked change: {str(e)}")
            return None

    def convert_to_html(self, docx_path: str) -> Dict[str, Any]:
        """
        Convert DOCX to HTML for Lexical editor import.

        This is a simple conversion. For production, consider using mammoth.js
        via subprocess or a Python HTML generation library.

        Args:
            docx_path: Path to the DOCX file

        Returns:
            Dictionary with HTML content and success status
        """
        try:
            doc = DocxDocument(docx_path)
            html_parts = ['<div>']

            for para in doc.paragraphs:
                if not para.text.strip():
                    continue

                # Determine heading level
                style_name = para.style.name if para.style else "Normal"
                if "Heading 1" in style_name:
                    tag = "h1"
                elif "Heading 2" in style_name:
                    tag = "h2"
                elif "Heading 3" in style_name:
                    tag = "h3"
                elif "Heading 4" in style_name:
                    tag = "h4"
                elif "Heading 5" in style_name:
                    tag = "h5"
                elif "Heading 6" in style_name:
                    tag = "h6"
                else:
                    tag = "p"

                # Build paragraph HTML with inline formatting
                para_html = f"<{tag}>"
                for run in para.runs:
                    text = run.text
                    if run.bold:
                        text = f"<strong>{text}</strong>"
                    if run.italic:
                        text = f"<em>{text}</em>"
                    if run.underline:
                        text = f"<u>{text}</u>"
                    para_html += text
                para_html += f"</{tag}>"

                html_parts.append(para_html)

            # Add tables
            for table in doc.tables:
                html_parts.append('<table border="1">')
                for row in table.rows:
                    html_parts.append('<tr>')
                    for cell in row.cells:
                        html_parts.append(f'<td>{cell.text}</td>')
                    html_parts.append('</tr>')
                html_parts.append('</table>')

            html_parts.append('</div>')
            html_content = '\n'.join(html_parts)

            return {
                "success": True,
                "html": html_content
            }

        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }

    def get_document_metadata(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract metadata from DOCX file (title, author, creation date, etc.).

        Args:
            docx_path: Path to the DOCX file

        Returns:
            Dictionary with document metadata
        """
        try:
            doc = DocxDocument(docx_path)
            core_props = doc.core_properties

            return {
                "title": core_props.title or "Untitled Document",
                "author": core_props.author or "Unknown",
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "subject": core_props.subject,
                "keywords": core_props.keywords
            }

        except Exception as e:
            return {
                "title": "Untitled Document",
                "author": "Unknown",
                "error": str(e)
            }
